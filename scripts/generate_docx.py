"""Generate a sample Word document to verify the harness works."""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


def generate_sample_docx(output_dir: str = "output") -> str:
    """Generate a sample document with various content types.

    Returns the path to the generated file.
    """
    doc = Document()

    # --- Page setup (A4) ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)

    # --- Header ---
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "Document Harness Verification"
    hp.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # --- Document properties ---
    doc.core_properties.author = "Claude Code Harness"
    doc.core_properties.title = "Sample Document"

    # --- Title ---
    title = doc.add_heading("Sample Document", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Introduction ---
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This document verifies the Word document generation capability "
        "of the Claude Code document harness."
    )

    # --- Formatted text ---
    doc.add_heading("2. Formatting Examples", level=1)
    doc.add_heading("2.1 Inline Formatting", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Bold text, ")
    run.bold = True
    run = p.add_run("italic text, ")
    run.italic = True
    run = p.add_run("colored text, ")
    run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    p.add_run("and normal text.")

    # --- Bullet points ---
    doc.add_heading("2.2 Lists", level=2)
    for item in [
        "Bullet point one",
        "Bullet point two",
        "Bullet point three",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    for item in [
        "Numbered item one",
        "Numbered item two",
        "Numbered item three",
    ]:
        doc.add_paragraph(item, style="List Number")

    # --- Table ---
    doc.add_heading("3. Data Table", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Shading Accent 1"
    for i, h in enumerate(["Item", "Current", "Target", "Progress"]):
        table.rows[0].cells[i].text = h
    data = [
        ("Revenue", "$100K", "$150K", "67%"),
        ("Users", "200", "300", "67%"),
        ("Retention", "85%", "90%", "94%"),
    ]
    for row_data in data:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val

    # --- Page break + Appendix ---
    doc.add_page_break()
    doc.add_heading("Appendix", level=1)
    doc.add_paragraph("Supplementary data would go here.")

    # Save
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = os.path.join(output_dir, f"document_{timestamp}.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_sample_docx()
