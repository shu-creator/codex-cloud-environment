"""Tests for Word document generation."""

import os
import tempfile

from docx import Document
from docx.shared import Cm

from scripts.generate_docx import generate_sample_docx


def test_generate_sample_docx_creates_file():
    """Verify that generate_sample_docx produces a valid .docx file."""
    with tempfile.TemporaryDirectory() as tmpdir:
        path = generate_sample_docx(output_dir=tmpdir)
        assert os.path.exists(path)
        assert path.endswith(".docx")


def test_generated_docx_has_content():
    """Verify the generated document is not empty."""
    with tempfile.TemporaryDirectory() as tmpdir:
        path = generate_sample_docx(output_dir=tmpdir)
        doc = Document(path)
        assert len(doc.paragraphs) > 0


def test_generated_docx_has_table():
    """Verify the generated document contains a table."""
    with tempfile.TemporaryDirectory() as tmpdir:
        path = generate_sample_docx(output_dir=tmpdir)
        doc = Document(path)
        assert len(doc.tables) >= 1


def test_generated_docx_is_a4():
    """Verify the generated document uses A4 page size."""
    with tempfile.TemporaryDirectory() as tmpdir:
        path = generate_sample_docx(output_dir=tmpdir)
        doc = Document(path)
        section = doc.sections[0]
        # A4: 21cm x 29.7cm (with small rounding tolerance)
        assert abs(section.page_width - Cm(21.0)) < Cm(0.1)
        assert abs(section.page_height - Cm(29.7)) < Cm(0.1)


def test_generated_docx_has_header():
    """Verify the generated document has a header."""
    with tempfile.TemporaryDirectory() as tmpdir:
        path = generate_sample_docx(output_dir=tmpdir)
        doc = Document(path)
        header = doc.sections[0].header
        assert len(header.paragraphs) > 0
        assert header.paragraphs[0].text != ""
